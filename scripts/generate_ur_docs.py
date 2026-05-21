#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate printable Word docs for user research learning."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT_DIR = Path("/Users/dadaizhi/Desktop")


def set_doc_style(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    style.font.size = Pt(11)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "PingFang SC"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(10)
        r2.font.name = "PingFang SC"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    doc.add_paragraph()


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.name = "PingFang SC"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.name = "PingFang SC"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.name = "PingFang SC"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def add_p(doc, text, bold=False):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.bold = bold
        r.font.name = "PingFang SC"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.name = "PingFang SC"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def add_page_break(doc):
    doc.add_page_break()


def build_learning_doc():
    doc = Document()
    set_doc_style(doc)
    add_title(doc, "用户研究 · 学习手册", "打印版 | 黄素知 | 配合九牧等智能硬件用研岗位")

    add_p(doc, "说明：本手册整理自岗位 JD 与现有项目经验差距，侧重「问卷交叉分析、访谈提炼、可用性测试、电商评论洞察」。设计岗转用研不必先学复杂统计，Excel + 清晰表达即可。")

    add_h1(doc, "一、你需要补什么（对照九牧 JD）")
    add_bullets(doc, [
        "问卷：清洗 → 描述统计 → 交叉分析（透视表）→ 3～5 条发现 → 建议",
        "访谈：原话摘录 → 打标签 → 主题归类（亲和图）→ 按角色/场景交叉对比",
        "可用性：任务成功率、耗时、错误 → 严重度 0～4 → 优化建议",
        "电商评论：抽样 50～100 条 → 归类标签 → 高频痛点 → 1 条产品建议",
        "入户/情境（了解即可）：观察环境 + 真实任务 + 深访",
        "暂不深学：SPSS 全套、复杂回归、抽样公式推导",
    ])

    add_h1(doc, "二、2 天学习计划（可勾选）")
    add_h2(doc, "第 1 天：问卷交叉 + 优先级 + UME 复述")
    add_bullets(doc, [
        "□ 上午：看 1～2 个「Excel 数据透视表 问卷交叉」短视频（各 15 分钟内）",
        "□ 上午：用 30 份假数据做 1 张「角色 × 满意度」交叉表",
        "□ 下午：学痛点优先级（影响面 × 严重程度，或简化 ICE）",
        "□ 下午：把 UME「100+ 问题分层」写成 3 分钟口述稿（见另一份文档）",
        "□ 晚上：读 30 条智能马桶/淋浴电商差评，手工贴 8～10 个标签",
    ])
    add_h2(doc, "第 2 天：访谈提炼 + 厨卫练手报告 + 模拟面试")
    add_bullets(doc, [
        "□ 上午：学亲和图（摘录→标签→主题），用 5 段原话练习",
        "□ 下午：按《智能厨卫评论洞察练手》大纲完成 8 页迷你报告",
        "□ 下午：整理可用性测试报告模板（背景/目标/方法/发现/建议）",
        "□ 晚上：自答 6 道模拟面试题（每题 2 分钟内，见本手册第六节）",
    ])

    add_page_break(doc)
    add_h1(doc, "三、问卷交叉分析（实操 Cheat Sheet）")
    add_p(doc, "研究问题：不同____（角色/是否买过/年龄档）的人，在____（某题/某痛点）上是否有差异？")
    add_h3(doc, "步骤（Excel）")
    add_bullets(doc, [
        "1. 导出问卷 CSV，先清洗：删无效卷、作答时间过短、逻辑矛盾项",
        "2. 插入 → 数据透视表",
        "3. 行 = 分组变量（如：角色-行政/员工/管理员）",
        "4. 列 = 题目选项或「是否同意/满意度档位」",
        "5. 值 = 计数（人数）",
        "6. 计算占比：某格人数 ÷ 该行合计（或 ÷ 总样本，二选一，全文统一）",
        "7. 找「某一格占比明显偏高/偏低」→ 写成发现",
        "8. 每条发现 = 数据描述 + 业务解释 + 1 条可执行建议",
    ])
    add_h3(doc, "面试一句话")
    add_p(doc, "「交叉分析就是看不同用户群在同一问题上的差异；我用 Excel 透视表看占比，再结合业务判断优先级，不是拍脑袋。」")

    add_h3(doc, "描述统计（必会）")
    add_bullets(doc, [
        "频数、百分比：某选项有多少人选",
        "均值：李克特 1～5 分题的平均分（仅作参考，小样本慎吹）",
        "开放题：不硬量化，做主题归类后数「提及次数」",
    ])

    add_h3(doc, "优先级矩阵（UME 100+ 问题可用）")
    add_bullets(doc, [
        "影响面：多少人遇到 / 多少业务线受影响",
        "严重程度：是否阻断任务、是否引发投诉/安全风险",
        "实现成本：研发改动大还是小（可与产品一起估）",
        "写法：高影响 + 高严重 → P0；其余 P1/P2",
    ])

    add_page_break(doc)
    add_h1(doc, "四、访谈：从原话到洞察")
    add_h3(doc, "四步流程")
    add_bullets(doc, [
        "1. 摘录：把原话按「一句一贴纸」拆开（Word/便签/飞书文档均可）",
        "2. 标签：任务 / 情绪 / 痛点 / 工具 / 期望（5 类即可）",
        "3. 主题：相似标签合并为 3～7 个主题（Theme）",
        "4. 交叉（定性版）：表格列=用户类型，行=主题，格子里写「提及次数 + 典型原话」",
    ])
    add_h3(doc, "深访提纲常用结构（30～45 分钟）")
    add_bullets(doc, [
        "暖场：说明目的、保密、录音同意",
        "背景：角色、使用频率、环境",
        "任务走查：最近一次完成某任务的全过程",
        "痛点：哪里卡、怎么绕、影响什么",
        "期望：理想状态、愿不愿付费/换品牌",
        "收尾：还有什么没说",
    ])

    add_h1(doc, "五、可用性测试报告骨架")
    add_bullets(doc, [
        "1. 背景与目标（测什么、为什么现在测）",
        "2. 方法（5～8 人、任务列表、环境、时长）",
        "3. 指标：任务成功率、平均完成时间、错误次数",
        "4. 发现：按任务列问题，标严重度 0～4（0=无，4=阻断）",
        "5. 建议：每条问题对应 1 条改法 + 负责人/版本（若已知）",
        "6. 附录：任务表、访谈/测试纪要",
    ])

    add_page_break(doc)
    add_h1(doc, "六、模拟面试 6 题（自答勾选）")
    questions = [
        ("介绍一个你独立推动的研究项目。", "主答 UME：背景→方法→样本→3 发现→落地。"),
        ("问卷样本多少？怎么招募？无效卷怎么处理？", "如实说样本量级；招募渠道；清洗规则。"),
        ("怎么做交叉分析？举一个「不同角色有差异」的例子。", "透视表 + 占比 + 业务解释。"),
        ("100+ 问题怎么排优先级？", "影响面×严重程度×成本；举例 1 个 P0。"),
        ("研究结论产品没采纳怎么办？", "记录原因；缩小范围再验证；找可落地切片。"),
        ("没做过厨卫/入户，怎么快速上手？", "方法可迁移；已做评论练手报告；愿意跟学入户。"),
    ]
    for i, (q, hint) in enumerate(questions, 1):
        add_p(doc, f"□ 题 {i}：{q}", bold=True)
        add_p(doc, f"   提示：{hint}")
        doc.add_paragraph()

    add_h1(doc, "七、「研究体系」是什么（遇到快消/甲方 JD 时）")
    add_p(doc, "指公司内「怎么做研究、怎么沉淀、怎么驱动业务」的一整套机制，常包括：")
    add_bullets(doc, [
        "研究类型与节奏（概念测试/可用性/满意度/竞品/复盘）",
        "流程 SOP（立项→方案→招募→执行→报告→落地跟踪）",
        "模板（提纲、问卷、测试表、报告结构、严重度标准）",
        "样本库与招募规范",
        "洞察库（画像、痛点库、竞品档案）",
        "跨部门固定协作触点（评审会、洞察宣讲）",
        "第三方调研与预算管理（部分公司）",
    ])
    add_p(doc, "你可对标经历：两万里「访谈要点 + 问题识别标准 + 跨部门闭环」= 研究体系的局部；面试勿说成「完整快消研究体系从 0 搭建」。")

    add_h1(doc, "八、推荐学习资源（关键词搜索即可）")
    add_bullets(doc, [
        "B 站：问卷 交叉分析 Excel 数据透视表",
        "B 站：亲和图 affinity diagram 访谈整理",
        "书籍（选读章节）：《启示录》需求、《点石成金》可用性",
        "练习：腾讯问卷导出 CSV → 自己造 30 条也可练手",
    ])

    add_p(doc, "— 完 —", bold=True)
    return doc


def build_kitchen_outline_doc():
    doc = Document()
    set_doc_style(doc)
    add_title(doc, "《智能厨卫评论洞察练手》", "8 页大纲 | 面试展示用 | 填写说明见每页底部")

    add_p(doc, "目标：用公开信息完成一次「迷你用研」，证明你能做电商评论归类与洞察提炼。不必真实入户。建议总篇幅 6～10 页 PPT 或 Word，打印友好。")

    pages = [
        (
            "第 1 页｜封面",
            [
                "报告标题：智能厨卫用户评论洞察（练手）",
                "副标题：基于电商平台公开评论的初步分析",
                "姓名 / 日期",
                "声明：本报告为学习方法练手，数据来自公开评论，仅供个人学习展示",
            ],
            "【你要填】选品类：智能马桶 / 智能淋浴 / 浴室柜（三选一写清）",
        ),
        (
            "第 2 页｜研究背景与目标",
            [
                "背景：智能厨卫购买决策链长（比价、安装、售后），线上评论影响转化",
                "业务问题（自拟 1 句）：例如「首次购买用户最常吐槽什么？」",
                "研究目标（2～3 条）：",
                "  · 了解差评/问大家中的高频痛点类别",
                "  · 对比 2～3 个竞品品牌的评论差异",
                "  · 提出 1～2 条可讨论的产品/体验优化方向",
                "不涉及：入户观察（本练手不做，可写「后续可深化」）",
            ],
            "【你要填】为什么选这个品类（1～2 句，可写投九牧相关）",
        ),
        (
            "第 3 页｜方法与数据来源",
            [
                "方法：桌面研究 + 公开评论定性归类（轻量）",
                "平台：京东 / 天猫（写清你用的平台）",
                "样本：",
                "  · 品牌 A / B / C 各抽差评约 15～20 条，合计约 50～80 条",
                "  · 可选：「问大家」10 条补充",
                "时间范围：近 6 个月或默认排序前 N 条",
                "局限（必写，显专业）：",
                "  · 样本非随机，不能代表全国用户",
                "  · 评论用户偏负面表达，需结合销量与星级综合看",
            ],
            "【你要填】实际品牌名、每条平台、抽样日期",
        ),
        (
            "第 4 页｜评论编码表（核心页）",
            [
                "建一张表，列：序号 | 品牌 | 星级 | 原文摘要（20 字内）| 一级标签 | 二级标签",
                "一级标签建议（可增删）：",
                "  · 安装与施工  · 功能与冲水/出水  · 噪音与震动",
                "  · 材质与做工  · 智能/App/遥控  · 售后与物流  · 价格与促销",
                "二级标签示例：",
                "  · 安装：师傅不专业 / 尺寸不合 / 漏水",
                "  · 智能：连不上 App / 感应失灵 / 耗电",
                "本页展示 8～12 条代表性条目即可，完整表放附录",
            ],
            "【你要填】至少编码 50 条；统计各一级标签「条数」",
        ),
        (
            "第 5 页｜频次统计与交叉（简化版）",
            [
                "表 1：一级标签 × 条数 × 占比（饼图或条形图可选）",
                "表 2（简化交叉）：品牌 × 一级标签（哪品牌在「安装」上提及更多）",
                "写法示例：",
                "  「在 N 条差评中，安装类占 X%，高于功能类 Y%」",
                "  「品牌 A 的安装相关提及占比高于品牌 B」",
                "Excel：用 COUNTIF 或数据透视表即可",
            ],
            "【你要填】真实数字；没有数据先空着，打印后手写补",
        ),
        (
            "第 6 页｜典型用户声音（Quote）",
            [
                "选 3～5 条最有代表性的原话（可打码昵称）",
                "每条格式：引语 + 标签 + 你的 1 句解读",
                "示例结构：",
                "  「装完第二天就漏水…」→ 安装与施工 → 反映安装质检与售后响应重要",
                "可选：粗分人群（新房装修 / 换新 / 酒店工装）— 若评论里能看出来",
            ],
            "【你要填】必须从真实评论复制，勿编造",
        ),
        (
            "第 7 页｜洞察与建议",
            [
                "洞察 3 条（每条 3 行）：现象（数据）→ 原因假设 → 对产品/服务的启示",
                "建议 2 条（可执行、不过度承诺）：",
                "  · 产品侧：例「安装指引可视化」「首次配对向导」",
                "  · 服务侧：例「安装后 24h 回访」「差评关键词监控看板」",
                "风险：建议需研发/供应链配合的，标「需进一步验证」",
            ],
            "【你要填】洞察必须挂钩第 5 页数据",
        ),
        (
            "第 8 页｜总结与后续研究",
            [
                "3 行总结本练手收获",
                "后续若入职可深化：",
                "  · 问卷验证痛点优先级",
                "  · 入户观察真实浴室动线",
                "  · 可用性测试 App/面板交互",
                "  · 与产品/研发的需求评审节奏",
                "附录（可选加页）：完整编码表、竞品功能对比表",
            ],
            "【你要填】面试时口述 3 分钟：背景→方法→1 个数据→2 个洞察→1 个建议",
        ),
    ]

    for title, bullets, note in pages:
        add_h2(doc, title)
        add_bullets(doc, bullets)
        add_p(doc, note)
        doc.add_paragraph()

    add_p(doc, "— 大纲完，请按页填写后打印 —", bold=True)
    return doc


def build_ume_script_doc():
    doc = Document()
    set_doc_style(doc)
    add_title(doc, "UME 协同平台改版 · 面试口述稿", "约 3～5 分钟 | 黄素知 | 打印后可在空白处补充样本量与数据")

    add_p(doc, "使用说明：这是「用户研究视角」的讲述顺序，不是背稿。括号内为可补充项，请按你真实经历修改。")

    add_h1(doc, "〇、开场（20 秒）")
    add_p(doc, "「我介绍一个在企业协同软件改版里，从研究驱动到方案落地的主项目——亿联 UME。产品要从聊天工具升级成协同办公平台，我主动发起并主导了这次改版里的用户研究工作，并参与后续的交互框架与多端方案。」")

    add_h1(doc, "一、背景与为什么要做研究（30 秒）")
    add_bullets(doc, [
        "业务背景：业务线扩展，旧版信息架构扛不住新功能，多端体验不一致",
        "不能只靠设计师主观判断，需要证据支撑「改什么、先改什么」",
        "研究目标：① 找核心痛点 ② 区分不同角色诉求 ③ 为 IA 重构和优先级提供依据",
    ])

    add_h1(doc, "二、我负责什么（20 秒）")
    add_p(doc, "「研究侧我负责定目标、选方法、组织执行、归纳发现、做优先级，并参与方案评审；交付物包括问题清单、优先级矩阵、研究报告要点，以及后续交互框架输入。覆盖 PC 桌面端与移动端核心场景。」")

    add_h1(doc, "三、方法与过程（60～90 秒）— 重点")
    add_h3(doc, "3.1 竞品研究")
    add_p(doc, "「先看竞品：对标主流协同产品，梳理功能模式与体验模式，明确差距与可借鉴点。」")

    add_h3(doc, "3.2 用户访谈")
    add_p(doc, "「做深度访谈，覆盖（补充：___ 人，角色包括 ___）。关注真实任务：沟通、检索、会议预约、视频会议等。」")

    add_h3(doc, "3.3 问卷调研")
    add_p(doc, "「用问卷扩大样本，验证访谈里高频问题是否普遍。回收后先做清洗（无效卷、作答时间过短等），再统计。」")

    add_h3(doc, "3.4 交叉与优先级（可接 Excel 透视表）")
    add_p(doc, "「例如按角色看『会议相关痛点』占比是否更高——这就是交叉分析思路。最终沉淀 100+ 体验问题，并按影响面、严重程度、改造成本做分层，形成 P0/P1/P2，支撑排期。」")

    add_h3(doc, "3.5 方案收敛")
    add_p(doc, "「基于研究输出 3 套方向，组织内外部评审达成共识；同时梳理近 300 个需求点，搭建可扩展信息架构，打通核心流程。」")

    add_h1(doc, "四、3 条代表性发现（60 秒）— 请改成你的原话")
    add_p(doc, "发现 1（沟通/检索）：", bold=True)
    add_p(doc, "「用户在 ___ 场景下频繁 ___，导致 ___。（补充 1 句典型原话或数据）」")
    add_p(doc, "发现 2（会议/视频）：", bold=True)
    add_p(doc, "「会中信息干扰明显，后续支撑了会议模块『人像聚焦』等策略。」")
    add_p(doc, "发现 3（平台扩展）：", bold=True)
    add_p(doc, "「旧 IA 难以支撑平台化，需要从信息架构层重构，而不是只改界面。」")

    add_h1(doc, "五、落地与结果（30 秒）")
    add_bullets(doc, [
        "研究结论直接进入排期与方案，不是写完报告就结束",
        "新版上线后，用户反馈与可用性感知优于旧版（补充：若有问卷复测/NPS/工单数据可写）",
        "与 UC 生态、硬件产品线打通，研究也服务了多端一致性",
    ])

    add_h1(doc, "六、反思与诚实边界（20 秒，加分）")
    add_p(doc, "「这是设计岗位里研究比重很高的项目，不是公司专职用研 title；定量上我以问卷占比和优先级为主，没有做成大型市场研究。若入职九牧，我会把同样流程用在厨卫场景，并补强入户与电商评论类研究。」")

    add_h1(doc, "七、面试官可能追问 → 简短答法")
    faq = [
        ("样本各多少？", "访谈（  ）人；问卷（  ）份有效；竞品（  ）款。——请填真实数字。"),
        ("100+ 问题怎么来的？", "访谈纪要 + 问卷开放题 + 走查 + 内部反馈，合并去重后打标签。"),
        ("交叉分析具体例子？", "透视表：行=角色，列=是否遇到某痛点，看占比差异。"),
        ("产品没采纳某条建议怎么办？", "记录原因；拆小步验证；或转为下一版本议题。"),
        ("你和交互设计的边界？", "研究定『改什么』，交互定『怎么改』；本项目两者我都参与，但决策先有证据。"),
    ]
    for q, a in faq:
        add_p(doc, f"Q：{q}", bold=True)
        add_p(doc, f"A：{a}")

    add_page_break(doc)
    add_h1(doc, "附：口述时间轴（可剪到 3 分钟）")
    add_bullets(doc, [
        "0:00-0:20  开场 + 项目名",
        "0:20-0:50  背景 + 研究目标",
        "0:50-2:20  方法（竞品/访谈/问卷/100+优先级）",
        "2:20-3:20  2～3 条发现 + 1 条落地",
        "3:20-3:40  结果",
        "3:40-4:00  反思边界（可选）",
    ])

    add_p(doc, "— 打印后在「补充」栏填写样本量、原话、数据 —", bold=True)
    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    files = [
        (build_learning_doc(), "01-用户研究学习手册.docx"),
        (build_kitchen_outline_doc(), "02-智能厨卫评论洞察练手-8页大纲.docx"),
        (build_ume_script_doc(), "03-UME项目面试口述稿.docx"),
    ]
    for doc, name in files:
        path = OUT_DIR / name
        doc.save(path)
        print(f"Saved: {path}")


if __name__ == "__main__":
    main()
